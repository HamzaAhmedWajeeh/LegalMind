"""
core/ingestion/parser.py
========================
Document parser — extracts raw text from uploaded files.

Supports:
  - PDF  : pdfplumber (fast, structured) with pytesseract OCR fallback
             for scanned/image-only PDFs
  - DOCX : python-docx
  - TXT  : plain read

Design Pattern: Strategy Pattern (partial) — the parser selects the
correct extraction strategy based on file type/content automatically.
The chunker (chunker.py) applies the full Strategy Pattern on top of
whatever raw text this module produces.

Returns a ParsedDocument dataclass consumed by the chunker.
"""

import hashlib
import io
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
# pytesseract imported lazily in _ocr_page() to avoid pandas/numpy conflict
import structlog
from PIL import Image
from docx import Document as DocxDocument

logger = structlog.get_logger(__name__)

# Minimum character count per page before we consider it a scanned page
# and fall back to OCR.
_OCR_FALLBACK_THRESHOLD = 50


# ──────────────────────────────────────────────────────────────────
# Data transfer object produced by the parser
# ──────────────────────────────────────────────────────────────────
@dataclass
class ParsedDocument:
    """
    Raw extraction result passed downstream to the chunker.

    Attributes:
        filename      : Original file name
        file_hash     : SHA-256 of the raw bytes (deduplication key)
        raw_text      : Full extracted text, pages joined by form-feed \\f
        page_count    : Number of pages (PDFs) or sections (DOCX)
        file_type     : 'pdf' | 'docx' | 'txt'
        ocr_used      : True if any page required OCR fallback
        per_page_text : List of per-page strings (useful for page-level metadata)
        metadata      : Any extra key-value pairs extracted during parsing
    """
    filename: str
    file_hash: str
    raw_text: str
    page_count: int
    file_type: str
    ocr_used: bool = False
    per_page_text: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


# ──────────────────────────────────────────────────────────────────
# Public entry point
# ──────────────────────────────────────────────────────────────────
def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Parse a document from raw bytes.

    Args:
        file_bytes : Raw file content
        filename   : Original filename (used to determine type)

    Returns:
        ParsedDocument with extracted text and metadata

    Raises:
        ValueError  : If the file type is unsupported
        RuntimeError: If extraction fails entirely
    """
    file_hash = _sha256(file_bytes)
    extension = Path(filename).suffix.lower()

    logger.info("Parsing document", filename=filename, size_bytes=len(file_bytes))

    if extension == ".pdf":
        return _parse_pdf(file_bytes, filename, file_hash)
    elif extension == ".docx":
        return _parse_docx(file_bytes, filename, file_hash)
    elif extension in (".txt", ".md"):
        return _parse_text(file_bytes, filename, file_hash)
    else:
        raise ValueError(
            f"Unsupported file type '{extension}'. "
            "Supported: .pdf, .docx, .txt, .md"
        )


# ──────────────────────────────────────────────────────────────────
# PDF parser
# ──────────────────────────────────────────────────────────────────
def _parse_pdf(
    file_bytes: bytes,
    filename: str,
    file_hash: str,
) -> ParsedDocument:
    """
    Extract text from PDF using pdfplumber.
    Falls back to pytesseract OCR for pages with insufficient text
    (i.e., scanned image pages).
    """
    per_page_text: list[str] = []
    ocr_used = False

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""

                # If the page has very little text, it's likely a scanned image
                if len(text.strip()) < _OCR_FALLBACK_THRESHOLD:
                    logger.debug(
                        "Page has insufficient text — attempting OCR",
                        filename=filename,
                        page=page_num,
                        text_len=len(text.strip()),
                    )
                    ocr_text = _ocr_page(page)
                    if ocr_text:
                        text = ocr_text
                        ocr_used = True

                per_page_text.append(text)

    except Exception as exc:
        raise RuntimeError(
            f"Failed to parse PDF '{filename}': {exc}"
        ) from exc

    raw_text = "\f".join(per_page_text)   # form-feed as page separator

    logger.info(
        "PDF parsed",
        filename=filename,
        pages=page_count,
        total_chars=len(raw_text),
        ocr_used=ocr_used,
    )

    return ParsedDocument(
        filename=filename,
        file_hash=file_hash,
        raw_text=raw_text,
        page_count=page_count,
        file_type="pdf",
        ocr_used=ocr_used,
        per_page_text=per_page_text,
        metadata={"page_count": page_count, "ocr_used": ocr_used},
    )


def _ocr_page(page) -> str:
    """
    Render a pdfplumber page to an image and run pytesseract OCR.
    Returns empty string if OCR is unavailable or fails.
    pytesseract is imported lazily here to avoid its pandas dependency
    conflicting with deepeval's numpy at module load time.
    """
    try:
        try:
            import pytesseract  # noqa: PLC0415
        except ImportError:
            logger.debug("pytesseract not installed — OCR unavailable")
            return ""
        # Render at 300 DPI for good OCR accuracy
        pil_image = page.to_image(resolution=300).original
        text = pytesseract.image_to_string(pil_image, lang="eng")
        return text.strip()
    except Exception as exc:
        logger.warning("OCR failed for page", error=str(exc))
        return ""


# ──────────────────────────────────────────────────────────────────
# DOCX parser
# ──────────────────────────────────────────────────────────────────
def _parse_docx(
    file_bytes: bytes,
    filename: str,
    file_hash: str,
) -> ParsedDocument:
    """Extract text from a .docx file paragraph by paragraph."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n\n".join(paragraphs)
    except Exception as exc:
        raise RuntimeError(
            f"Failed to parse DOCX '{filename}': {exc}"
        ) from exc

    logger.info(
        "DOCX parsed",
        filename=filename,
        paragraphs=len(paragraphs),
        total_chars=len(raw_text),
    )

    return ParsedDocument(
        filename=filename,
        file_hash=file_hash,
        raw_text=raw_text,
        page_count=len(paragraphs),    # paragraphs as proxy for "pages"
        file_type="docx",
        per_page_text=paragraphs,
        metadata={"paragraph_count": len(paragraphs)},
    )


# ──────────────────────────────────────────────────────────────────
# Plain text parser
# ──────────────────────────────────────────────────────────────────
def _parse_text(
    file_bytes: bytes,
    filename: str,
    file_hash: str,
) -> ParsedDocument:
    """Decode a plain text file."""
    try:
        raw_text = file_bytes.decode("utf-8", errors="replace")
    except Exception as exc:
        raise RuntimeError(
            f"Failed to parse text file '{filename}': {exc}"
        ) from exc

    logger.info("Text file parsed", filename=filename, total_chars=len(raw_text))

    return ParsedDocument(
        filename=filename,
        file_hash=file_hash,
        raw_text=raw_text,
        page_count=1,
        file_type="txt",
        per_page_text=[raw_text],
    )


# ──────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────
def _sha256(data: bytes) -> str:
    """Return hex SHA-256 digest of raw bytes."""
    return hashlib.sha256(data).hexdigest()
